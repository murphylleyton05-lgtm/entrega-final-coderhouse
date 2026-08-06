"""
Genera el documento anexo con el registro de prompts de IA (.docx y .pdf).
El contenido sale de construir.ANALISIS_IA, para que la hoja Analisis_IA del
libro y el documento no puedan desincronizarse.
"""
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

from construir import ANALISIS_IA

NAVY = RGBColor(0x1F, 0x38, 0x64)
GRAY = RGBColor(0x6B, 0x70, 0x76)
AMBER = RGBColor(0xC1, 0x7D, 0x1F)


def _sombrear(par, hexcolor):
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def _p(doc, texto="", *, size=10.5, bold=False, italic=False, color=None,
       space_after=6, space_before=0, align=None, left=0.0):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    if left:
        par.paragraph_format.left_indent = Cm(left)
    if align:
        par.alignment = align
    run = par.add_run(texto)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return par


def construir_docx(destino):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Cm(2)
        sec.left_margin = sec.right_margin = Cm(2.2)
    est = doc.styles["Normal"]
    est.font.name = "Arial"
    est.font.size = Pt(10.5)

    _p(doc, "Registro de prompts de IA", size=20, bold=True, color=NAVY,
       space_after=2)
    _p(doc, "Validacion de hallazgos del Dashboard Final - TechStore",
       size=12, color=GRAY, space_after=2)
    _p(doc, "Lleyton Murphy  ·  Coderhouse, Data Analyst  ·  Modulo 9: capa "
            "de visualizacion", size=9.5, color=GRAY, space_after=14)

    _p(doc, "Como se uso la IA", size=13, bold=True, color=NAVY,
       space_before=6, space_after=4)
    _p(doc, "La IA no se uso para interpretar los graficos, porque eso es "
            "precisamente lo que genera deuda cognitiva: si le pido \"que dice "
            "este grafico\", dejo de analizar yo. Se uso como auditor: primero "
            "escribi las conclusiones a mano mirando el tablero, despues le "
            "pedi que las atacara buscando sesgos, errores de logica y "
            "variables omitidas, y recien despues verifique cada objecion "
            "contra los datos y corregi el diseno.")
    _p(doc, "El flujo, en tres pasos:", space_after=2)
    for t in ("1. Analisis humano: escribo las conclusiones sin ayuda.",
              "2. Contraste con IA: le doy contexto y conclusiones, y le pido "
              "sesgos y explicaciones alternativas, no un resumen.",
              "3. Verificacion y refinamiento: contrasto cada objecion contra "
              "los datos y cambio el tablero donde la objecion tenia razon."):
        _p(doc, t, left=0.6, space_after=3)
    _p(doc, "Herramienta: ChatGPT (GPT-5). Las respuestas estan resumidas en su "
            "sustancia: se transcribe la objecion, no el texto completo.",
       size=9, italic=True, color=GRAY, space_after=12)

    _p(doc, "El resultado en una linea", size=13, bold=True, color=NAVY,
       space_before=6, space_after=4)
    _p(doc, "De las tres conclusiones que escribi, una se confirmo con un "
            "matiz, una se refuto en su atribucion causal y una estaba "
            "comparada contra el periodo equivocado. Dos elementos del tablero "
            "-el grafico de cascada y la comparativa 2024-2025 del grafico "
            "mensual- existen por esas objeciones: no estaban en el diseno "
            "original.", space_after=14)

    etiquetas = [("propia", "Conclusion propia (paso 1)"),
                 ("prompt", "Prompt enviado (paso 2)"),
                 ("objecion", "Objecion de la IA"),
                 ("verificacion", "Verificacion sobre los datos (paso 3)"),
                 ("refinamiento", "Cambio aplicado al tablero")]

    for i, bloque in enumerate(ANALISIS_IA, start=1):
        doc.add_page_break() if i > 1 else None
        cab = _p(doc, "  " + bloque["titulo"], size=12.5, bold=True,
                 color=NAVY, space_before=4, space_after=8)
        _sombrear(cab, "EDF0F5")
        for clave, etiqueta in etiquetas:
            _p(doc, etiqueta, size=10, bold=True,
               color=AMBER if clave == "objecion" else NAVY,
               space_before=6, space_after=2)
            texto = bloque[clave]
            if clave == "prompt":
                par = _p(doc, texto, size=10, italic=True, left=0.5)
                _sombrear(par, "F7F8FA")
            else:
                _p(doc, texto, left=0.5)

    doc.add_page_break()
    _p(doc, "Que quedo sin resolver", size=13, bold=True, color=NAVY,
       space_after=4)
    _p(doc, "La IA planteo dos objeciones que este tablero todavia no puede "
            "responder, y quedan anotadas como limitacion honesta y no como "
            "hallazgo:")
    for t in ("El dataset no tiene costo de adquisicion de cliente por canal. "
              "Sin ese dato no se puede afirmar que el e-commerce sea menos "
              "rentable que la tienda fisica: solo que su costo logistico por "
              "peso vendido es mayor.",
              "Tampoco hay devoluciones. En e-commerce la tasa de devolucion "
              "suele ser mayor que en tienda, asi que el margen real del canal "
              "podria ser todavia mas bajo que el que muestra el tablero."):
        _p(doc, "·  " + t, left=0.6, space_after=6)

    _p(doc, "Nota sobre los datos", size=13, bold=True, color=NAVY,
       space_before=10, space_after=4)
    _p(doc, "El dataset es simulado y determinista (semilla fija). Extiende el "
            "modelo Ventas_Tech_DB del checkpoint de SQL -mismas categorias, "
            "productos y ciudades- a 24 meses, 5 regiones y 2 canales, para "
            "poder analizar margen, estacionalidad y evolucion interanual. "
            "Todos los supuestos y metas estan documentados en la hoja "
            "Parametros del libro.")

    doc.save(destino)
    return destino


def a_pdf(docx_path, outdir):
    subprocess.run(
        ["soffice", "--headless", "--norestore",
         "-env:UserInstallation=file:///tmp/lo_doc",
         "--convert-to", "pdf", "--outdir", outdir, docx_path],
        check=True, capture_output=True, timeout=600)
    return os.path.join(
        outdir, os.path.basename(docx_path).replace(".docx", ".pdf"))


if __name__ == "__main__":
    destino = sys.argv[1]
    os.makedirs(os.path.dirname(destino), exist_ok=True)
    construir_docx(destino)
    print("docx:", destino)
    print("pdf :", a_pdf(destino, os.path.dirname(destino)))
